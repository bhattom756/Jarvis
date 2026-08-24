from __future__ import annotations

import logging
import mimetypes
import os
import smtplib
import ssl
import subprocess
import webbrowser
from datetime import timedelta
from email.message import EmailMessage
from pathlib import Path
from urllib.parse import quote_plus

try:
    from playwright.sync_api import BrowserContext, Playwright, sync_playwright
except ImportError:  # pragma: no cover - optional runtime dependency
    BrowserContext = Playwright = None  # type: ignore[assignment]
    sync_playwright = None  # type: ignore[assignment]

try:
    import pyautogui
except ImportError:  # pragma: no cover - optional runtime dependency
    pyautogui = None  # type: ignore[assignment]

try:
    import pygetwindow
except ImportError:  # pragma: no cover - optional runtime dependency
    pygetwindow = None  # type: ignore[assignment]

from app.config import Settings, settings as default_settings
from app.schemas import ActionResult, ConfirmationPayload, EmailDraftRequest, utc_now


logger = logging.getLogger(__name__)


class BrowserActionProvider:
    def __init__(self, settings: Settings | None = None) -> None:
        self.settings = settings or default_settings
        self._playwright: Playwright | None = None
        self._context: BrowserContext | None = None

    def search(self, query: str) -> ActionResult:
        return self.navigate(f"https://www.google.com/search?q={quote_plus(query)}", summary=f"Searched the web for '{query}'.")

    def navigate(self, url: str, summary: str | None = None) -> ActionResult:
        target = url if url.startswith(("http://", "https://")) else f"https://{url}"
        if not self.settings.enable_browser_actions:
            return ActionResult(ok=False, summary="Browser actions are disabled.")
        try:
            page = self._page()
            page.goto(target, wait_until="domcontentloaded", timeout=30_000)
            title = page.title()
            return ActionResult(ok=True, summary=summary or f"Opened {title or target}.", detail=page.url)
        except Exception as exc:
            logger.warning("Playwright navigation failed; opening default browser: %s", exc)
            try:
                webbrowser.open(target)
                return ActionResult(ok=True, summary=summary or f"Opened {target} in the default browser.", detail=target)
            except Exception as fallback_exc:
                return ActionResult(ok=False, summary="Unable to open browser.", detail=str(fallback_exc))

    def extract(self, url: str) -> ActionResult:
        try:
            page = self._page()
            page.goto(url, wait_until="domcontentloaded", timeout=30_000)
            text = page.locator("body").inner_text(timeout=10_000)
            excerpt = " ".join(text.split())[:2000]
            return ActionResult(ok=True, summary=f"Extracted page content from {page.title() or url}.", detail=excerpt)
        except Exception as exc:
            return ActionResult(ok=False, summary="Unable to extract page content.", detail=str(exc))

    def close(self) -> None:
        if self._context is not None:
            self._context.close()
        if self._playwright is not None:
            self._playwright.stop()
        self._context = None
        self._playwright = None

    def _page(self):
        if sync_playwright is None:
            raise RuntimeError("Playwright is not installed")
        if self._context is None:
            self._playwright = sync_playwright().start()
            self._context = self._playwright.chromium.launch_persistent_context(
                user_data_dir=str(self.settings.browser_session_dir),
                headless=self.settings.browser_headless,
            )
        return self._context.pages[0] if self._context.pages else self._context.new_page()


class WindowsActionProvider:
    APP_ALIASES = {
        "whatsapp": "whatsapp:",
        "word": "winword",
        "ms word": "winword",
        "microsoft word": "winword",
        "excel": "excel",
        "powerpoint": "powerpnt",
        "ppt": "powerpnt",
        "vscode": "code",
        "visual studio code": "code",
        "code": "code",
        "chrome": "chrome",
        "edge": "msedge",
        "browser": "msedge",
        "notepad": "notepad",
        "calculator": "calc",
        "calc": "calc",
        "file explorer": "explorer",
        "explorer": "explorer",
        "spotify": "spotify",
        "settings": "ms-settings:",
    }

    def open_application(self, target: str) -> ActionResult:
        cleaned_target = target.lower().strip()
        executable = self.APP_ALIASES.get(cleaned_target, target)
        try:
            if executable.endswith(":"):
                subprocess.Popen(f'start "" "{executable}"', shell=True)
            else:
                try:
                    subprocess.Popen([executable], shell=False)
                except OSError:
                    subprocess.Popen(f'start "" "{executable}"', shell=True)
            return ActionResult(ok=True, summary=f"Launched {target} successfully, Boss.")
        except Exception as exc:
            try:
                os.startfile(target)  # type: ignore[attr-defined]
                return ActionResult(ok=True, summary=f"Launched {target} successfully, Boss.")
            except Exception as fallback_exc:
                return ActionResult(ok=False, summary=f"Unable to launch {target}.", detail=str(fallback_exc))

    def create_word_document(self, filename: str, title: str = "", content: str = "") -> ActionResult:
        try:
            from docx import Document
            doc = Document()
            if title:
                doc.add_heading(title, 0)
            if content:
                for paragraph in content.split("\n\n"):
                    if paragraph.strip():
                        doc.add_paragraph(paragraph.strip())
            else:
                doc.add_paragraph("Document created by FRIDAY.")

            target_name = filename if filename.endswith(".docx") else f"{filename}.docx"
            desktop_path = Path.home() / "Desktop" / target_name
            doc.save(str(desktop_path))
            return ActionResult(ok=True, summary=f"Created Word document '{target_name}' on your Desktop, Boss.", detail=str(desktop_path))
        except Exception as exc:
            logger.exception("Failed to create Word document")
            return ActionResult(ok=False, summary=f"Failed to create Word document: {exc}", detail=str(exc))

    def create_pdf_document(self, filename: str, title: str = "", content: str = "") -> ActionResult:
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas
            target_name = filename if filename.endswith(".pdf") else f"{filename}.pdf"
            desktop_path = Path.home() / "Desktop" / target_name

            c = canvas.Canvas(str(desktop_path), pagesize=letter)
            c.setFont("Helvetica-Bold", 16)
            c.drawString(72, 750, title or filename.replace(".pdf", ""))
            c.setFont("Helvetica", 11)
            y = 710
            lines = (content or "Document generated by FRIDAY.").split("\n")
            for line in lines:
                if y < 72:
                    c.showPage()
                    c.setFont("Helvetica", 11)
                    y = 750
                c.drawString(72, y, line[:95])
                y -= 18
            c.save()
            return ActionResult(ok=True, summary=f"Created PDF document '{target_name}' on your Desktop, Boss.", detail=str(desktop_path))
        except Exception as exc:
            logger.exception("Failed to create PDF document")
            return ActionResult(ok=False, summary=f"Failed to create PDF document: {exc}", detail=str(exc))

    def list_windows(self) -> ActionResult:
        if pygetwindow is None:
            return ActionResult(ok=False, summary="Window discovery requires pygetwindow.")
        titles = [title for title in pygetwindow.getAllTitles() if title.strip()]
        return ActionResult(ok=True, summary=f"Found {len(titles)} open windows.", detail="\n".join(titles[:30]))

    def focus_window(self, title: str) -> ActionResult:
        if pygetwindow is None:
            return ActionResult(ok=False, summary="Window switching requires pygetwindow.")
        matches = pygetwindow.getWindowsWithTitle(title)
        if not matches:
            return ActionResult(ok=False, summary=f"No window matched '{title}'.")
        window = matches[0]
        window.restore()
        window.activate()
        return ActionResult(ok=True, summary=f"Focused window '{window.title}'.")

    def type_text(self, text: str) -> ActionResult:
        if pyautogui is None:
            return ActionResult(ok=False, summary="Keyboard automation requires pyautogui.")
        pyautogui.write(text, interval=0.02)
        return ActionResult(ok=True, summary="Typed text into the active window.")


class EmailActionProvider:
    """SMTP sender used only after an explicit confirmation is resolved."""

    def __init__(self, settings: Settings) -> None:
        self.settings = settings

    def validate(self, draft: EmailDraftRequest) -> ActionResult:
        if not self._configured:
            return ActionResult(ok=False, summary="SMTP email sending is not configured.")
        try:
            attachment_names = [self._validate_attachment(Path(path)).name for path in draft.attachments]
        except ValueError as exc:
            return ActionResult(ok=False, summary="Email draft contains an invalid attachment.", detail=str(exc))
        attachment_note = f" Attachments: {', '.join(attachment_names)}." if attachment_names else ""
        return ActionResult(ok=True, summary=f"Send email to {', '.join(draft.recipients)} with subject '{draft.subject}'.{attachment_note}")

    def send(self, draft: EmailDraftRequest) -> ActionResult:
        validation = self.validate(draft)
        if not validation.ok:
            return validation
        message = EmailMessage()
        message["From"] = self.settings.email_from_address or self.settings.email_smtp_username or self.settings.email_imap_username
        message["To"] = ", ".join(draft.recipients)
        message["Subject"] = draft.subject
        message.set_content(draft.body)
        try:
            for path_str in draft.attachments:
                path = self._validate_attachment(Path(path_str))
                mime_type, _ = mimetypes.guess_type(path.name)
                major, minor = (mime_type or "application/octet-stream").split("/", 1)
                message.add_attachment(path.read_bytes(), maintype=major, subtype=minor, filename=path.name)
            context = ssl.create_default_context()
            if self.settings.email_smtp_use_ssl:
                with smtplib.SMTP_SSL(self.settings.email_smtp_host, self.settings.email_smtp_port, context=context, timeout=30) as client:
                    client.login(self._username, self._password)
                    client.send_message(message)
            else:
                with smtplib.SMTP(self.settings.email_smtp_host, self.settings.email_smtp_port, timeout=30) as client:
                    client.starttls(context=context)
                    client.login(self._username, self._password)
                    client.send_message(message)
            return ActionResult(ok=True, summary=f"Email sent to {', '.join(draft.recipients)}.", detail=validation.summary)
        except Exception as exc:
            logger.exception("SMTP email send failed")
            return ActionResult(ok=False, summary="Email was not sent.", detail=str(exc))

    @property
    def _username(self) -> str:
        return self.settings.email_smtp_username or self.settings.email_imap_username or ""

    @property
    def _password(self) -> str:
        return self.settings.email_smtp_password or self.settings.email_imap_password or ""

    @property
    def _configured(self) -> bool:
        return bool(self.settings.email_smtp_host and self._username and self._password and (self.settings.email_from_address or self._username))

    def _validate_attachment(self, path: Path) -> Path:
        resolved = path.expanduser().resolve()
        if not resolved.is_file():
            raise ValueError(f"Attachment does not exist or is not a file: {path}")
        if resolved.stat().st_size > self.settings.email_max_attachment_bytes:
            raise ValueError(f"Attachment exceeds {self.settings.email_max_attachment_bytes // 1_000_000} MB: {resolved.name}")
        return resolved


class WeatherActionProvider:
    def get_weather(self, location: str = "Bangalore") -> ActionResult:
        try:
            import httpx
            loc_clean = location.strip(" .?!,").title() or "Bengaluru"
            if loc_clean.lower() in ("bangalore", "bangalore city", "the bagel", "bagel"):
                loc_clean = "Bengaluru"

            geo_resp = httpx.get(
                "https://geocoding-api.open-meteo.com/v1/search",
                params={"name": loc_clean, "count": 5, "language": "en", "format": "json"},
                timeout=5.0,
            )
            geo_data = geo_resp.json()
            results = geo_data.get("results", [])
            if not results:
                loc_clean = "Bengaluru"
                geo_resp = httpx.get(
                    "https://geocoding-api.open-meteo.com/v1/search",
                    params={"name": loc_clean, "count": 5, "language": "en", "format": "json"},
                    timeout=5.0,
                )
                results = geo_resp.json().get("results", [])

            results.sort(key=lambda r: r.get("population", 0), reverse=True)
            result = results[0]
            lat, lon = result["latitude"], result["longitude"]
            city_name = result.get("name", loc_clean)
            country = result.get("country", "")

            weather_resp = httpx.get(
                "https://api.open-meteo.com/v1/forecast",
                params={
                    "latitude": lat,
                    "longitude": lon,
                    "current": "temperature_2m,relative_humidity_2m,apparent_temperature,precipitation,weather_code,wind_speed_10m",
                    "timezone": "auto",
                },
                timeout=5.0,
            )
            w_data = weather_resp.json().get("current", {})
            temp = w_data.get("temperature_2m", "N/A")
            feels_like = w_data.get("apparent_temperature", temp)
            humidity = w_data.get("relative_humidity_2m", "N/A")
            wind = w_data.get("wind_speed_10m", "N/A")
            w_code = w_data.get("weather_code", 0)

            conditions = {
                0: "Clear sky", 1: "Mainly clear", 2: "Partly cloudy", 3: "Overcast",
                45: "Foggy", 48: "Depositing rime fog", 51: "Light drizzle", 53: "Moderate drizzle",
                61: "Slight rain", 63: "Moderate rain", 65: "Heavy rain", 80: "Rain showers",
                95: "Thunderstorm"
            }
            condition_desc = conditions.get(w_code, "Partly cloudy")

            summary = f"Current weather in {city_name}, {country}: {temp}°C (feels like {feels_like}°C), {condition_desc}, Humidity: {humidity}%, Wind speed: {wind} km/h."
            return ActionResult(ok=True, summary=summary, detail=str(w_data))
        except Exception as exc:
            logger.warning("Weather fetch failed: %s", exc)
            return ActionResult(ok=False, summary=f"Could not retrieve weather data for {location}: {exc}")


class ActionEngine:
    def __init__(self, settings: Settings, browser_provider: BrowserActionProvider, windows_provider: WindowsActionProvider) -> None:
        self.settings = settings
        self.browser = browser_provider
        self.windows = windows_provider
        self.weather = WeatherActionProvider()

    def evaluate(self, utterance: str, approved: bool = False) -> ActionResult:
        if self._requires_confirmation(utterance) and not approved:
            return self._confirmation(utterance, "high_risk_action", "high", reversible=False)

        lower = utterance.lower().strip()

        if any(w in lower for w in ("weather", "temperature", "climate", "forecast", "how hot", "how cold")):
            loc = "Bangalore"
            for marker in (" in ", " for ", " at "):
                if marker in lower:
                    loc = utterance.split(marker, 1)[1].strip(" .?!")
                    break
            return self.weather.get_weather(loc)

        if any(phrase in lower for phrase in ("time right now", "what's the time", "what time is it", "current time", "today's date", "what date is it")):
            from datetime import datetime
            now_str = datetime.now().strftime("%I:%M %p, %A, %B %d, %Y")
            return ActionResult(ok=True, summary=f"Current local time and date: {now_str}.")

        if not self.settings.enable_windows_actions and any(lower.startswith(prefix) for prefix in ("open ", "launch ", "list windows", "switch to ", "focus ", "type ")):
            return ActionResult(ok=False, summary="Windows actions are disabled.")

        if lower.startswith(("create word document", "make word document", "create docx", "generate word document")):
            parts = utterance.split(" ", 3)
            filename = parts[3] if len(parts) > 3 else "Document"
            return self.windows.create_word_document(filename, title="Document", content=utterance)

        if lower.startswith(("create pdf", "make pdf", "generate pdf")):
            parts = utterance.split(" ", 2)
            filename = parts[2] if len(parts) > 2 else "Document"
            return self.windows.create_pdf_document(filename, title="Document", content=utterance)

        if lower.startswith(("open ", "launch ")):
            target = utterance.split(" ", 1)[1].strip()
            if target.startswith(("http://", "https://")) or ("." in target and not target.endswith((":", " app"))):
                return self.browser.navigate(target)
            return self.windows.open_application(target)
        if "search for" in lower:
            return self.browser.search(utterance.split("search for", 1)[1].strip())
        if lower.startswith("go to ") or lower.startswith("navigate to "):
            return self.browser.navigate(utterance.split(" ", 2)[-1].strip())
        if lower.startswith("extract ") or lower.startswith("read page "):
            return self.browser.extract(utterance.split(" ", 1)[1].strip())
        if "list windows" in lower or "open windows" in lower:
            return self.windows.list_windows()
        if lower.startswith("switch to ") or lower.startswith("focus "):
            return self.windows.focus_window(utterance.split(" ", 2)[-1].strip())
        if lower.startswith("type "):
            return self.windows.type_text(utterance.split(" ", 1)[1])
        if lower.startswith("create file "):
            return self._create_file(utterance.split("create file ", 1)[1].strip(), approved)
        return ActionResult(ok=True, summary="No direct action executed. Plan updated for response.")

    def close(self) -> None:
        self.browser.close()

    def _create_file(self, path_str: str, approved: bool) -> ActionResult:
        path = Path(path_str)
        if not self._is_approved_path(path) and not approved:
            return self._confirmation(f"Create file {path}", "file_write", "medium", reversible=True)
        path.parent.mkdir(parents=True, exist_ok=True)
        path.touch(exist_ok=True)
        return ActionResult(ok=True, summary=f"Created file at {path}.")

    def _confirmation(self, summary: str, category: str, risk_level: str, reversible: bool) -> ActionResult:
        return ActionResult(
            ok=False,
            summary="Confirmation required before execution.",
            confirmation=ConfirmationPayload(
                category=category,
                risk_level=risk_level,  # type: ignore[arg-type]
                summary=summary,
                reversible=reversible,
                expires_at=utc_now() + timedelta(minutes=5),
            ),
        )

    def _requires_confirmation(self, utterance: str) -> bool:
        risky_terms = ("delete", "send", "purchase", "shutdown", "format", "move all", "run script", "type ")
        return any(term in utterance.lower() for term in risky_terms)

    def _is_approved_path(self, path: Path) -> bool:
        resolved = path.resolve()
        for root in self.settings.approved_write_roots:
            try:
                resolved.relative_to(Path(root).resolve())
                return True
            except ValueError:
                continue
        return False
